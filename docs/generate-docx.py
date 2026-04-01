#!/usr/bin/env python3
"""
Generate Official AI — Pipeline Architecture DOCX
Professional document with visual state machine diagrams, flow boxes, and tables.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Style config ──
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica Neue'
font.size = Pt(10)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x0f, 0x0f, 0x23)
    hs.font.name = 'Helvetica Neue'

# Colors
INDIGO = RGBColor(0x63, 0x66, 0xf1)
VIOLET = RGBColor(0x8b, 0x5c, 0xf6)
CYAN = RGBColor(0x06, 0xb6, 0xd4)
EMERALD = RGBColor(0x10, 0xb9, 0x81)
AMBER = RGBColor(0xf5, 0x9e, 0x0b)
RED = RGBColor(0xef, 0x44, 0x44)
SLATE = RGBColor(0x64, 0x74, 0x8b)
DARK = RGBColor(0x0f, 0x17, 0x2a)
WHITE = RGBColor(0xff, 0xff, 0xff)

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="6366f1"):
    """Create a styled table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
        set_cell_bg(cell, header_color)

    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_bg(cell, "f1f5f9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table

def add_state_box(doc, label, description, color_hex="6366f1", width_pct=100):
    """Create a visual state box using a single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]

    # Top border color
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="4" w:color="e2e8f0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="e2e8f0"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="e2e8f0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    set_cell_bg(cell, "f8fafc")

    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16))

    p2 = cell.add_paragraph()
    run2 = p2.add_run(description)
    run2.font.size = Pt(9)
    run2.font.color.rgb = SLATE

def add_flow_row(doc, steps, colors=None):
    """Create a horizontal flow of state boxes with arrows between them."""
    if colors is None:
        colors = ["6366f1"] * len(steps)

    cols = len(steps) * 2 - 1  # steps + arrows
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in t.rows:
        for cell in row.cells:
            # Remove all borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)

    step_idx = 0
    for i in range(cols):
        cell = t.rows[0].cells[i]
        if i % 2 == 0:
            # State box
            label, desc = steps[step_idx]
            color = colors[step_idx]

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders2 = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="12" w:color="{color}"/>'
                f'  <w:left w:val="single" w:sz="4" w:color="cbd5e1"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:color="cbd5e1"/>'
                f'  <w:right w:val="single" w:sz="4" w:color="cbd5e1"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders2)
            set_cell_bg(cell, "f8fafc")

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))

            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(desc)
            run2.font.size = Pt(7)
            run2.font.color.rgb = SLATE

            step_idx += 1
        else:
            # Arrow
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("\u2192")
            run.font.size = Pt(16)
            run.font.color.rgb = SLATE

    doc.add_paragraph()

def add_connector_arrow(doc, label=""):
    """Add a downward arrow between sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label:
        run = p.add_run(f"\u2193  {label}  \u2193")
    else:
        run = p.add_run("\u2193")
    run.font.size = Pt(14)
    run.font.color.rgb = SLATE


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OFFICIAL AI")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = INDIGO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Pipeline Architecture & API Reference")
run2.font.size = Pt(18)
run2.font.color.rgb = SLATE

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Technical Documentation")
run3.font.size = Pt(12)
run3.font.color.rgb = SLATE

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("March 31, 2026  |  Mudita Studios")
run4.font.size = Pt(11)
run4.font.color.rgb = SLATE

doc.add_paragraph()
doc.add_paragraph()

# Version box
add_state_box(doc, "Document Version", "v2.0 — Updated with V2 Onboarding (4-step), Voice Clone, 7 video formats, complete API map, and state machine diagrams.", "6366f1")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Architecture Overview",
    "2. Video Generation Pipeline — State Machine",
    "3. Pipeline Steps (7-Step Breakdown)",
    "4. Onboarding Flow (V2)",
    "5. External Service Integration",
    "6. TTS Provider Fallback Chain",
    "7. Video Model Routing",
    "8. Shotstack Stitching & Timeline",
    "9. Video Formats & Cut Patterns",
    "10. API Route Reference",
    "11. Data Shapes & Interfaces",
    "12. Environment Variables",
    "13. File Reference",
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════

doc.add_heading("1. Architecture Overview", level=1)

p = doc.add_paragraph()
run = p.add_run("Official AI generates professional AI videos through a 7-step serverless pipeline. ")
run.font.size = Pt(10)
run = p.add_run("Every step is designed to run within Netlify's 26-second function timeout. Heavy work is delegated to external APIs (FAL.ai, Shotstack, Gemini) with async polling.")
run.font.size = Pt(10)

doc.add_paragraph()

# Design principles boxes
principles = [
    ("Serverless-Safe", "All heavy processing delegated to external APIs. No step exceeds 26s.", "6366f1"),
    ("Webhook + Polling", "FAL webhooks drive pipeline progression. Client polling is the failsafe.", "8b5cf6"),
    ("Audio-Driven Composition", "TTS duration determines video cut trim values. Audio leads, video follows.", "06b6d4"),
    ("Idempotent State Machine", "Multiple advance calls produce the same result. Safe for webhook + polling.", "10b981"),
    ("Per-Cut Audio Alignment", "Each video cut gets its own audio segment on Shotstack timeline. No bleed.", "f59e0b"),
]

for label, desc, color in principles:
    add_state_box(doc, label, desc, color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. STATE MACHINE
# ═══════════════════════════════════════════════════════════

doc.add_heading("2. Video Generation Pipeline — State Machine", level=1)

p = doc.add_paragraph("The video record's status field acts as a finite state machine. Each state has defined transitions and failure modes.")
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# Top-level state flow
doc.add_heading("Top-Level Status Lifecycle", level=2)

add_flow_row(doc,
    [
        ("DRAFT", "Video created\nComposition planned"),
        ("GENERATING", "Pipeline active\n7 substeps"),
        ("REVIEW", "Video delivered\nAwaiting approval"),
        ("APPROVED", "User approved\nReady to publish"),
        ("PUBLISHED", "Live on\nsocial platforms"),
    ],
    ["64748b", "6366f1", "f59e0b", "10b981", "06b6d4"]
)

# Failed state
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GENERATING can also transition to ")
run.font.size = Pt(9)
run = p.add_run("FAILED")
run.font.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RED
run = p.add_run(" on: step error, 5-minute timeout, or stitch failure")
run.font.size = Pt(9)

doc.add_paragraph()

# Pipeline substates
doc.add_heading("Pipeline Substates (within GENERATING)", level=2)

add_flow_row(doc,
    [
        ("expand", "Gemini\nscript"),
        ("tts", "Audio\nper-cut"),
        ("anchor", "Starting\nframe"),
        ("submit", "FAL\njobs"),
    ],
    ["6366f1", "8b5cf6", "06b6d4", "f59e0b"]
)

add_flow_row(doc,
    [
        ("poll_cuts", "FAL\npolling"),
        ("stitch", "Shotstack\nsubmit"),
        ("poll_stitch", "Shotstack\npolling"),
        ("DONE", "Video\ndelivered"),
    ],
    ["f59e0b", "10b981", "10b981", "06b6d4"]
)

doc.add_paragraph()

# Progress percentages
doc.add_heading("Progress Calculation", level=2)

add_styled_table(doc,
    ["Step", "Progress %", "External Service"],
    [
        ["queued", "0%", "—"],
        ["expand", "5%", "Gemini Flash"],
        ["tts", "15%", "FAL MiniMax / MiniMax / ElevenLabs"],
        ["anchor", "20%", "Gemini Image Gen"],
        ["submit_all_cuts", "25%", "FAL.ai (Kling/MiniMax/WAN)"],
        ["poll_all_cuts", "25–85%", "FAL.ai polling"],
        ["stitch", "90%", "Shotstack"],
        ["store", "95%", "Supabase Storage"],
        ["done", "100%", "—"],
    ],
    [1.5, 1.0, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Auto-timeout: ")
run.font.bold = True
run.font.size = Pt(9)
run = p.add_run("GENERATION_TIMEOUT_MS = 300,000ms (5 minutes). Videos in 'generating' status longer than this are automatically marked 'failed'.")
run.font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. PIPELINE STEPS
# ═══════════════════════════════════════════════════════════

doc.add_heading("3. Pipeline Steps — Detailed Breakdown", level=1)

steps = [
    ("Step 1: EXPAND", "6366f1", "Script Expansion via Gemini",
     "File: src/lib/pipeline/expand.ts\nExternal: Gemini Flash via content-planner.ts\n\nSingle unified API call to Gemini merges prompt expansion and per-cut planning. Each cut receives: prompt (scene description), script (dialogue/narration), targetDuration. Results stored in video.sourceReview.cuts[]."),

    ("Step 2: TTS", "8b5cf6", "Text-to-Speech Generation (3-Provider Fallback)",
     "File: src/lib/pipeline/tts.ts\nExternal: FAL MiniMax \u2192 MiniMax Standalone \u2192 ElevenLabs\n\nIterates each cut, generates TTS for its script text. Audio duration determines cut trim values (audio-driven composition). TTS failure is NON-FATAL \u2014 pipeline continues without audio for that cut. Each provider gets 2 retries with exponential backoff."),

    ("Step 3: ANCHOR", "06b6d4", "Starting Frame Resolution",
     "File: src/lib/pipeline/anchor.ts\nExternal: Gemini Image Generation API\n\nResolves or generates a starting frame (reference image) for character consistency across all cuts. Uses user photos + character sheet as input. Prompt: iPhone 15 Pro Max 24mm f/1.78, medium close-up, natural window light. Critical for lip-sync and testimonial formats."),

    ("Step 4: SUBMIT ALL CUTS", "f59e0b", "Parallel FAL Video Submission",
     "File: src/lib/pipeline/cut-submit-all.ts\nExternal: FAL.ai (Kling 2.6 / MiniMax Hailuo / WAN 2.1)\n\nSubmits ALL cuts in parallel to FAL.ai. Model selected per cut via model-router.ts routing table. Each submission includes: prompt, reference image, audio URL, duration target. Returns immediately with job IDs."),

    ("Step 5: POLL ALL CUTS", "f59e0b", "Parallel Job Polling",
     "File: src/lib/pipeline/cut-poll-all.ts\nExternal: FAL.ai status API\n\nUses Promise.allSettled() for parallel polling. Reduces sequential 3\u20136 minute processing to 60\u2013120 seconds. Webhook callbacks from FAL also trigger progression. Returns nextStep: 'poll_all_cuts' with retryAfter: 5s if still processing."),

    ("Step 6: STITCH", "10b981", "Shotstack Video Composition",
     "File: src/lib/pipeline/stitch-submit.ts\nExternal: Shotstack Cloud Video Editing API\n\nConstructs Shotstack timeline: Track 0 = video clips with cross-dissolve transitions. Track 1 = per-cut audio aligned to each video cut's timeline position. Output: mp4, 9:16 aspect ratio, 30fps, high quality."),

    ("Step 7: POLL STITCH", "10b981", "Final Video Delivery",
     "File: src/lib/pipeline/stitch-poll.ts\nExternal: Shotstack + Supabase Storage\n\nPolls Shotstack render job. On completion: downloads video, uploads to Supabase Storage as permanent URL. Updates video record: videoUrl, thumbnailUrl, status: 'review'. FALLBACK: If stitch fails, uses first completed cut as final video."),
]

for label, color, subtitle, desc in steps:
    add_state_box(doc, f"{label} \u2014 {subtitle}", desc, color)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. ONBOARDING FLOW
# ═══════════════════════════════════════════════════════════

doc.add_heading("4. Onboarding Flow (V2 \u2014 4 Steps)", level=1)

p = doc.add_paragraph("The V2 onboarding flow captures photo, generates AI twin, clones voice, then presents the paywall with a live video preview.")

doc.add_paragraph()

# Step boxes
onboarding_steps = [
    ("Step 1: Photo Capture", "6366f1",
     "Component: CameraCapture.tsx\nAPIs: POST /api/upload, POST /api/photos\nCapture: MediaDevices.getUserMedia (webcam) or file upload\nQuality gate: min 300\u00d7300px, min 50KB\nEvent: onboarding_photo_captured"),

    ("Step 2: AI Twin \u2014 Character Sheet", "8b5cf6",
     "Component: CharacterSheetReveal.tsx\nAPI: POST /api/character-sheet\nGeneration: Gemini Flash \u2014 3\u00d73 pose grid + 2\u00d73 360\u00b0 sheet in parallel\nReveal: Cinematic animation with confetti burst (~4\u201310s generation)\nUser selects a pose \u2192 stored as characterSheetUrl\nEvent: onboarding_character_selected"),

    ("Step 3: Voice Clone", "06b6d4",
     "Component: VoiceCapture.tsx\nAPI: POST /api/onboarding/voice\nCapture: MediaRecorder API + AnalyserNode waveform visualization\nQuality gate: minimum 5 seconds recording\nModes: prompt \u2192 recording (live waveform) \u2192 preview (playback + re-record + confirm)\n4 rotating sample scripts for the user to read\nSkip option: 'Skip for now \u2014 you can add your voice later'\nEvents: onboarding_voice_cloned OR onboarding_voice_skipped"),

    ("Step 4: Paywall \u2014 Go Live", "6366f1",
     "Component: PaywallStep.tsx\nAPIs: POST /api/onboarding/preview-video (background), POST /api/stripe/checkout\nVideo preview: Shows generating spinner \u2192 play button when ready \u2192 tap to play/pause\nPricing: $79/mo after 7-day free trial\nFeatures: 6 bullet points, social proof (2 testimonials), ROI callout\nPaths: Start free trial \u2192 Stripe Checkout \u2192 Dashboard | Skip \u2192 Dashboard\nEvents: onboarding_paywall_viewed, onboarding_trial_started OR onboarding_skipped"),
]

for label, color, desc in onboarding_steps:
    add_state_box(doc, label, desc, color)
    if label != onboarding_steps[-1][0]:
        add_connector_arrow(doc)
    else:
        doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. EXTERNAL SERVICES
# ═══════════════════════════════════════════════════════════

doc.add_heading("5. External Service Integration", level=1)

services = [
    ("FAL.ai \u2014 Video Generation & Primary TTS", "6366f1",
     "Base URL: https://fal.run/fal-ai/{model}\nAuth: Authorization: Key {FAL_API_KEY}\nWebhook: POST /api/webhooks/fal (no auth)\n\nModels:\n  \u2022 kling-video/v2.5/standard \u2014 Primary video (best for people)\n  \u2022 minimax-video/video-01-live \u2014 Fallback video (best for scenes)\n  \u2022 wan-video/v2.1/1080p \u2014 Second fallback (general)\n  \u2022 minimax/speech-02-hd \u2014 Primary TTS"),

    ("Shotstack \u2014 Video Stitching", "10b981",
     "Base URL: https://api.shotstack.io/edit/{env}/render\nAuth: x-api-key: {SHOTSTACK_API_KEY}\nEnv: stage (sandbox) or v1 (production)\n\nRetry: 2 retries, 2s exponential backoff\nPolling: 3s \u2192 5s \u2192 8s \u2192 10s progressive backoff (max 5 min)\nResolution: sd / hd / 1080\nStatus: queued \u2192 processing \u2192 rendering \u2192 finalizing \u2192 completed/done/failed"),

    ("Google Gemini \u2014 Script & Image Generation", "f59e0b",
     "Base URL: https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent\nAuth: ?key={GOOGLE_AI_STUDIO_KEY}\n\nScript: Gemini Flash via content-planner.ts\nImages: nano-banana-pro-preview, gemini-2.5-flash-image, gemini-3-pro-image-preview, gemini-3.1-flash-image-preview\nConfig: responseModalities: ['image', 'text'], temperature: 0.6"),

    ("MiniMax Standalone \u2014 TTS (2nd Fallback)", "8b5cf6",
     "Base URL: https://api.minimax.chat/v1/t2a_v2?GroupId={groupId}\nAuth: Authorization: Bearer {MINIMAX_API_KEY}\nModel: speech-02-hd\nConfig: sample_rate 32000, bitrate 128000, format mp3, speed 1.0, pitch 0"),

    ("ElevenLabs \u2014 TTS (3rd Fallback)", "a78bfa",
     "Base URL: https://api.elevenlabs.io/v1/text-to-speech/{voiceId}\nAuth: xi-api-key: {ELEVENLABS_API_KEY}\nModel: eleven_multilingual_v2\nConfig: stability 0.5, similarity_boost 0.75, style 0.3, use_speaker_boost true\nDefault voice: 21m00Tcm4TlvDq8ikWAM"),

    ("Supabase \u2014 Database & Storage", "06b6d4",
     "Database: PostgreSQL via Prisma ORM (DATABASE_URL, DIRECT_URL)\nStorage buckets:\n  \u2022 photos/ \u2014 User photos\n  \u2022 starting-frames/{userId}/ \u2014 Character reference images\n  \u2022 videos/{userId}/{videoId}.mp4 \u2014 Final rendered videos\n  \u2022 voice-samples/ \u2014 Voice clone audio"),

    ("PostBridge \u2014 Social Publishing", "64748b",
     "Auth: POST_BRIDGE_API_KEY\nFlow: Upload video \u2192 Create post with caption \u2192 Store schedule record\nPlatforms: Instagram, TikTok, YouTube, LinkedIn, Facebook"),

    ("Stripe \u2014 Billing", "6366f1",
     "Checkout: POST /api/stripe/checkout creates session\nWebhook: POST /api/stripe/webhook verifies + processes events\nPlans: Starter ($79/mo, 7-day trial)\nKeys: STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET, NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY"),
]

for label, color, desc in services:
    add_state_box(doc, label, desc, color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. TTS FALLBACK CHAIN
# ═══════════════════════════════════════════════════════════

doc.add_heading("6. TTS Provider Fallback Chain", level=1)

p = doc.add_paragraph("Each provider gets 2 retries with exponential backoff before the next provider is tried. TTS failure is always non-fatal \u2014 the pipeline continues without audio.")

doc.add_paragraph()

add_flow_row(doc,
    [
        ("FAL MiniMax", "speech-02-hd\n2 retries"),
        ("MiniMax", "Standalone\n2 retries"),
        ("ElevenLabs", "multilingual_v2\n2 retries"),
        ("SKIP", "Non-fatal\nNo audio"),
    ],
    ["6366f1", "8b5cf6", "a78bfa", "64748b"]
)

doc.add_paragraph()

add_styled_table(doc,
    ["Provider", "Model", "Endpoint", "Auth Env Var"],
    [
        ["FAL MiniMax", "speech-02-hd", "fal.run/fal-ai/minimax/speech-02-hd", "FAL_API_KEY"],
        ["MiniMax Standalone", "speech-02-hd", "api.minimax.chat/v1/t2a_v2", "MINIMAX_API_KEY"],
        ["ElevenLabs", "eleven_multilingual_v2", "api.elevenlabs.io/v1/text-to-speech/{voiceId}", "ELEVENLABS_API_KEY"],
    ],
    [1.5, 1.5, 2.5, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Audio duration estimation: ")
run.font.bold = True
run.font.size = Pt(9)
run = p.add_run("text.split(/\\s+/).length / 2.5  (words per second)")
run.font.size = Pt(9)
run.font.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. VIDEO MODEL ROUTING
# ═══════════════════════════════════════════════════════════

doc.add_heading("7. Video Model Routing", level=1)

p = doc.add_paragraph("The model-router selects the primary video generation model based on cut type, with automatic fallback if the primary model fails.")

doc.add_paragraph()

doc.add_heading("Model Registry", level=2)

add_styled_table(doc,
    ["Internal ID", "FAL Endpoint", "Resolution", "Best For"],
    [
        ["kling-2.6", "fal-ai/kling-video/v2.5/standard", "1080p", "People / talking head"],
        ["minimax-hailuo", "fal-ai/minimax-video/video-01-live", "1080p", "Scenes / environments"],
        ["wan-2.1", "fal-ai/wan-video/v2.1/1080p", "1080p", "General fallback"],
    ],
    [1.5, 2.5, 1.0, 2.0]
)

doc.add_paragraph()

doc.add_heading("Routing Table", level=2)

add_styled_table(doc,
    ["Cut Type", "Primary", "Fallback 1", "Fallback 2"],
    [
        ["talking_head", "Kling 2.6", "MiniMax Hailuo", "WAN 2.1"],
        ["testimonial", "Kling 2.6", "MiniMax Hailuo", "WAN 2.1"],
        ["educational", "Kling 2.6", "MiniMax Hailuo", "WAN 2.1"],
        ["hook", "Kling 2.6", "MiniMax Hailuo", "WAN 2.1"],
        ["cta", "Kling 2.6", "MiniMax Hailuo", "WAN 2.1"],
        ["property_tour", "MiniMax Hailuo", "WAN 2.1", "Kling 2.6"],
        ["behind_scenes", "MiniMax Hailuo", "WAN 2.1", "Kling 2.6"],
    ],
    [1.5, 1.5, 1.5, 1.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 8. SHOTSTACK
# ═══════════════════════════════════════════════════════════

doc.add_heading("8. Shotstack Stitching & Timeline Construction", level=1)

p = doc.add_paragraph("Shotstack replaces FFmpeg for serverless-compatible video composition. The buildTimeline() function constructs a multi-track JSON payload.")

doc.add_paragraph()

doc.add_heading("Timeline Layout", level=2)

# Visual timeline
t = doc.add_table(rows=3, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = 'Table Grid'

# Header row
labels = ["", "Cut 1 (hook)", "Cut 2 (main)", "Cut 3 (point)", "Cut 4 (cta)"]
for i, l in enumerate(labels):
    cell = t.rows[0].cells[i]
    cell.text = l
    set_cell_bg(cell, "1e293b")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE

# Track 0 - Video
t.rows[1].cells[0].text = "Track 0\n(Video)"
set_cell_bg(t.rows[1].cells[0], "312e81")
for p in t.rows[1].cells[0].paragraphs:
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.bold = True

for i in range(1, 5):
    cell = t.rows[1].cells[i]
    cell.text = f"video clip {i}\n+ cross-dissolve"
    set_cell_bg(cell, "4338ca")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE

# Track 1 - Audio
t.rows[2].cells[0].text = "Track 1\n(Audio)"
set_cell_bg(t.rows[2].cells[0], "065f46")
for p in t.rows[2].cells[0].paragraphs:
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.bold = True

for i in range(1, 5):
    cell = t.rows[2].cells[i]
    cell.text = f"audio {i}\naligned to cut {i}"
    set_cell_bg(cell, "059669")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE

doc.add_paragraph()

add_state_box(doc, "Key Design: Per-Cut Audio Alignment",
    "Each audio segment starts exactly when its video cut starts.\nDuration = Math.min(audioDurationMs/1000, cutTrimTo) to prevent audio bleeding into next cut.\nThis solves sync issues vs. a single audio blob overlay approach.",
    "10b981")

doc.add_paragraph()

doc.add_heading("Output Configuration", level=2)

add_styled_table(doc,
    ["Parameter", "Value"],
    [
        ["Format", "mp4"],
        ["Resolution", "hd (configurable: sd / hd / 1080)"],
        ["Aspect Ratio", "9:16 (vertical)"],
        ["FPS", "30"],
        ["Quality", "high"],
        ["Background", "#000000"],
    ],
    [2.0, 4.0]
)

doc.add_paragraph()

doc.add_heading("Shotstack Status Flow", level=2)

add_flow_row(doc,
    [
        ("queued", "Job\nreceived"),
        ("processing", "Parsing\ntimeline"),
        ("rendering", "Compositing\nvideo"),
        ("completed", "Ready to\ndownload"),
    ],
    ["64748b", "f59e0b", "6366f1", "10b981"]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 9. VIDEO FORMATS
# ═══════════════════════════════════════════════════════════

doc.add_heading("9. Video Formats & Cut Patterns", level=1)

p = doc.add_paragraph("Seven defined formats in src/lib/video-compositor.ts. Each format specifies a sequence of typed cuts with target durations.")

doc.add_paragraph()

formats = [
    ("talking_head_15", "15s", [("hook", "3s"), ("main", "8s"), ("cta", "4s")]),
    ("testimonial_15", "15s", [("hook", "3s"), ("testimonial", "8s"), ("cta", "4s")]),
    ("testimonial_20", "20s", [("hook", "4s"), ("testimonial_a", "6s"), ("testimonial_b", "6s"), ("cta", "4s")]),
    ("educational_30", "30s", [("hook", "4s"), ("point_1", "7s"), ("point_2", "7s"), ("point_3", "7s"), ("cta", "5s")]),
    ("quick_tip_8", "8s", [("hook", "2s"), ("tip", "4s"), ("cta", "2s")]),
    ("property_tour_30", "30s", [("intro", "4s"), ("exterior", "6s"), ("interior_1", "6s"), ("interior_2", "6s"), ("feature", "5s"), ("cta", "3s")]),
    ("behind_scenes_20", "20s", [("hook", "3s"), ("scene_1", "6s"), ("scene_2", "6s"), ("reveal", "5s")]),
]

for fmt_name, total, cuts in formats:
    doc.add_heading(f"{fmt_name} ({total})", level=3)
    add_styled_table(doc,
        ["Cut #", "Type", "Duration"],
        [[str(i+1), c[0], c[1]] for i, c in enumerate(cuts)],
        [0.8, 2.0, 1.0]
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 10. API ROUTES
# ═══════════════════════════════════════════════════════════

doc.add_heading("10. API Route Reference", level=1)

sections = [
    ("Video Generation", "6366f1", [
        ["POST", "/api/generate", "Create video record, plan composition"],
        ["POST", "/api/generate/process", "Execute pipeline step {videoId, step, cutIndex?}"],
        ["POST", "/api/generate/advance", "Webhook/polling step progression (idempotent)"],
        ["GET", "/api/generate/status", "Progress polling with % calculation"],
        ["POST", "/api/generate/batch", "Batch create 3/5/7 from industry templates"],
        ["POST", "/api/generate/retry", "Retry from last failed step"],
    ]),
    ("Content & Media", "8b5cf6", [
        ["GET", "/api/videos", "List videos (excludes individual cuts)"],
        ["POST", "/api/upload", "Multipart upload (photo/voice/video)"],
        ["GET/POST", "/api/photos", "List / create photo records"],
        ["DELETE", "/api/photos/[id]", "Delete a photo"],
        ["GET/POST", "/api/character-sheet", "Generate or list character sheets"],
    ]),
    ("Onboarding", "06b6d4", [
        ["POST", "/api/onboarding/voice", "Upload voice sample for cloning"],
        ["POST", "/api/onboarding/preview-video", "Trigger preview video generation"],
        ["POST", "/api/onboarding/complete", "Mark onboarding finished"],
    ]),
    ("Publishing & Social", "10b981", [
        ["POST", "/api/publish", "Publish via PostBridge {videoId, platforms, caption}"],
        ["GET", "/api/social/accounts", "List connected accounts (tokens excluded)"],
    ]),
    ("Billing & Analytics", "f59e0b", [
        ["POST", "/api/stripe/checkout", "Create Stripe checkout session"],
        ["POST", "/api/stripe/webhook", "Stripe webhook handler"],
        ["GET", "/api/usage", "Plan usage metrics"],
        ["POST", "/api/events", "Track lifecycle events"],
    ]),
    ("Webhooks & Admin", "64748b", [
        ["POST", "/api/webhooks/fal", "FAL.ai job completion (no auth)"],
        ["GET/POST", "/api/admin/reset-stuck", "Count / reset stuck videos"],
        ["GET", "/api/admin/pipeline-log", "Pipeline event timeline"],
    ]),
]

for title, color, routes in sections:
    doc.add_heading(title, level=2)
    add_styled_table(doc,
        ["Method", "Route", "Purpose"],
        routes,
        [1.0, 2.5, 3.5],
        header_color=color
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 11. DATA SHAPES
# ═══════════════════════════════════════════════════════════

doc.add_heading("11. Data Shapes & Interfaces", level=1)

doc.add_heading("Video Record (Prisma)", level=2)

add_styled_table(doc,
    ["Field", "Type", "Notes"],
    [
        ["id", "String (cuid)", "Primary key"],
        ["userId", "String", "Foreign key to User"],
        ["title", "String?", "Optional"],
        ["script", "String?", "Raw input script"],
        ["model", "String?", "FAL model used"],
        ["contentType", "String?", "Format: talking_head_15, etc."],
        ["status", "String", "generating | review | approved | published | failed"],
        ["videoUrl", "String?", "Final video in Supabase"],
        ["thumbnailUrl", "String?", "Video thumbnail"],
        ["duration", "Float?", "Seconds"],
        ["sourceReview", "Json?", "Pipeline metadata (see below)"],
        ["photoId", "String?", "FK to Photo used"],
        ["voiceId", "String?", "FK to VoiceSample used"],
    ],
    [1.5, 1.5, 4.0]
)

doc.add_paragraph()

doc.add_heading("sourceReview Pipeline Metadata", level=2)

add_styled_table(doc,
    ["Field", "Type", "Description"],
    [
        ["pipelineStep", "string", "Current step name (expand, tts, anchor, etc.)"],
        ["pipelineCut", "number", "Progress counter during parallel cuts"],
        ["cuts[]", "Array", "Per-cut data: type, prompt, script, targetDuration, audio, videoUrl, status"],
        ["cuts[].audio.url", "string", "TTS audio URL for this cut"],
        ["cuts[].audio.duration", "number", "Audio duration in ms"],
        ["cuts[].audio.provider", "string", "fal-minimax | minimax | elevenlabs"],
        ["cutJobs{}", "Record", "Keyed by cut index: jobId, resultUrl"],
        ["cutThumbnailUrl", "string?", "Frame grab from first cut"],
        ["stitchJobId", "string?", "Shotstack render job ID"],
        ["error", "string?", "Error message on failure"],
    ],
    [2.0, 1.5, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 12. ENV VARS
# ═══════════════════════════════════════════════════════════

doc.add_heading("12. Environment Variables", level=1)

doc.add_heading("Required", level=2)

add_styled_table(doc,
    ["Variable", "Service", "Notes"],
    [
        ["DATABASE_URL", "Supabase PostgreSQL", "Connection pooling URL"],
        ["DIRECT_URL", "Supabase PostgreSQL", "Direct connection (migrations)"],
        ["NEXT_PUBLIC_SUPABASE_URL", "Supabase", "Public project URL"],
        ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Supabase", "Public anon key"],
        ["GOOGLE_AI_STUDIO_KEY", "Gemini API", "Script expansion + image gen"],
        ["FAL_API_KEY", "FAL.ai", "Video generation + TTS"],
        ["SHOTSTACK_API_KEY", "Shotstack", "Video stitching"],
        ["SHOTSTACK_ENV", "Shotstack", "stage (sandbox) or v1 (production)"],
    ],
    [2.5, 1.5, 3.0]
)

doc.add_paragraph()

doc.add_heading("TTS Providers (Fallback)", level=2)

add_styled_table(doc,
    ["Variable", "Service"],
    [
        ["MINIMAX_API_KEY", "MiniMax standalone TTS (2nd fallback)"],
        ["MINIMAX_GROUP_ID", "Required for MiniMax standalone API"],
        ["ELEVENLABS_API_KEY", "ElevenLabs TTS (3rd fallback)"],
    ],
    [2.5, 4.5]
)

doc.add_paragraph()

doc.add_heading("Social & Billing", level=2)

add_styled_table(doc,
    ["Variable", "Service"],
    [
        ["POST_BRIDGE_API_KEY", "PostBridge social distribution"],
        ["STRIPE_SECRET_KEY", "Stripe server-side billing"],
        ["STRIPE_WEBHOOK_SECRET", "Stripe webhook verification"],
        ["NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY", "Stripe client-side"],
        ["INSTAGRAM_CLIENT_ID / SECRET", "Instagram OAuth"],
        ["YOUTUBE_CLIENT_ID / SECRET", "YouTube OAuth"],
        ["TIKTOK_CLIENT_ID / SECRET", "TikTok OAuth"],
        ["LINKEDIN_CLIENT_ID / SECRET", "LinkedIn OAuth"],
        ["FACEBOOK_CLIENT_ID / SECRET", "Facebook OAuth"],
    ],
    [3.0, 4.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 13. FILE REFERENCE
# ═══════════════════════════════════════════════════════════

doc.add_heading("13. File Reference", level=1)

doc.add_heading("Pipeline Core", level=2)

add_styled_table(doc,
    ["File", "Purpose"],
    [
        ["src/lib/pipeline/orchestrator.ts", "Step router / state machine"],
        ["src/lib/pipeline/expand.ts", "Gemini script expansion"],
        ["src/lib/pipeline/tts.ts", "Per-cut TTS generation"],
        ["src/lib/pipeline/anchor.ts", "Starting frame resolution"],
        ["src/lib/pipeline/cut-submit-all.ts", "Parallel FAL submission"],
        ["src/lib/pipeline/cut-poll-all.ts", "Parallel FAL polling"],
        ["src/lib/pipeline/stitch-submit.ts", "Shotstack submission"],
        ["src/lib/pipeline/stitch-poll.ts", "Shotstack polling + persistence"],
        ["src/lib/pipeline/model-router.ts", "Cut type \u2192 model routing"],
        ["src/lib/pipeline/scene-bible.ts", "Environment generation via Gemini"],
    ],
    [3.5, 3.5]
)

doc.add_paragraph()

doc.add_heading("Service Clients", level=2)

add_styled_table(doc,
    ["File", "Purpose"],
    [
        ["src/lib/generate.ts", "FAL model registry (8 models)"],
        ["src/lib/video-stitcher.ts", "Shotstack API client"],
        ["src/lib/voice-engine.ts", "TTS with 3-provider fallback"],
        ["src/lib/video-compositor.ts", "7 format definitions"],
        ["src/lib/character-sheet.ts", "Pose + 360\u00b0 sheet generation"],
        ["src/lib/starting-frame.ts", "Anchor image generation"],
        ["src/lib/content-planner.ts", "Gemini script planning"],
    ],
    [3.5, 3.5]
)

doc.add_paragraph()

doc.add_heading("API Routes", level=2)

add_styled_table(doc,
    ["File", "Purpose"],
    [
        ["src/app/api/generate/route.ts", "Video creation"],
        ["src/app/api/generate/process/route.ts", "Pipeline step executor"],
        ["src/app/api/generate/advance/route.ts", "Step progression"],
        ["src/app/api/generate/status/route.ts", "Progress polling"],
        ["src/app/api/generate/batch/route.ts", "Batch creation"],
        ["src/app/api/webhooks/fal/route.ts", "FAL webhook handler"],
        ["src/app/api/upload/route.ts", "File upload"],
        ["src/app/api/publish/route.ts", "Social publishing"],
        ["src/app/api/onboarding/voice/route.ts", "Voice clone upload"],
        ["src/app/api/onboarding/preview-video/route.ts", "Preview video trigger"],
        ["src/app/api/stripe/checkout/route.ts", "Billing"],
        ["src/app/api/stripe/webhook/route.ts", "Stripe webhooks"],
    ],
    [3.5, 3.5]
)


# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Official-AI-Pipeline-Architecture.docx")
doc.save(output_path)
print(f"Generated: {output_path}")
